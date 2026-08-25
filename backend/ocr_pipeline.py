"""OCR and approval-note generation for synthetic or scanned inspection reports."""
from __future__ import annotations

import re
from pathlib import Path
from typing import Any

FIELD_LABELS = ("Equipment ID", "Inspection Date", "Inspector", "Finding", "Severity")


def extract_fields(raw_text: str) -> dict[str, str]:
    """Extract the five controlled labels from OCR text."""
    fields: dict[str, str] = {}
    for index, label in enumerate(FIELD_LABELS):
        next_labels = "|".join(re.escape(item) for item in FIELD_LABELS[index + 1 :])
        end = rf"(?=\s*(?:{next_labels})\s*:|\Z)" if next_labels else r"\Z"
        match = re.search(rf"{re.escape(label)}\s*:\s*(.+?){end}", raw_text, flags=re.IGNORECASE | re.DOTALL)
        fields[label] = " ".join(match.group(1).split()) if match else "Not detected"
    return fields


def ocr_and_extract(image_path: str | Path) -> dict[str, str]:
    """Run local Tesseract OCR, then extract controlled inspection fields."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError as error:
        raise RuntimeError("OCR requires pytesseract and Pillow. Install backend/requirements.txt.") from error

    try:
        with Image.open(image_path) as image:
            image.verify()
        with Image.open(image_path) as image:
            if image.width > 8000 or image.height > 8000:
                raise RuntimeError("Inspection image dimensions must be 8000x8000 pixels or smaller.")
            raw_text = pytesseract.image_to_string(image)
    except pytesseract.pytesseract.TesseractNotFoundError as error:
        raise RuntimeError("Tesseract executable is not installed or is not on PATH.") from error
    except (OSError, ValueError) as error:
        raise RuntimeError("The uploaded file is not a valid inspection image.") from error
    return extract_fields(raw_text)


def recommendation_for(severity: str) -> str:
    normalized = severity.lower()
    if normalized.strip() == "high":
        return "Stop or isolate affected equipment and escalate for immediate engineering review."
    if normalized.strip() == "medium":
        return "Create a priority maintenance work order and inspect during the next controlled window."
    if normalized.strip() == "low":
        return "Continue monitored operation and include the finding in the next planned inspection."
    return "Human review required because severity was missing or uncertain."


def write_approval_note(fields: dict[str, str], output_path: str | Path) -> Path:
    """Write extracted fields and a severity-based recommendation to a real DOCX."""
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError("DOCX generation requires python-docx. Install backend/requirements.txt.") from error

    destination = Path(output_path)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_heading("MRPL Inspection Approval Note", level=0)
    document.add_paragraph("Generated locally by the Sovereign AI Workbench")
    table = document.add_table(rows=1, cols=2)
    table.style = "Light Shading Accent 1"
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    for label in FIELD_LABELS:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = fields.get(label, "Not detected")
    document.add_heading("Finding", level=1)
    document.add_paragraph(fields.get("Finding", "Not detected"))
    document.add_heading("Recommendation", level=1)
    document.add_paragraph(recommendation_for(fields.get("Severity", "")))
    document.save(destination)
    return destination
